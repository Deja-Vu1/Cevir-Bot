import discord
from discord.ext import commands
from googletrans import Translator
import googletrans
from discord.ext import tasks
from bs4 import BeautifulSoup
import requests
from docx import Document
from docx.shared import Inches
import pytz
from pytz import timezone
from datetime import datetime
# This bot doing anything :D, Made by love <3

PERFIX = ">"
bot = commands.Bot(command_prefix=PERFIX)
bot.remove_command('help')
chk = 0
adminid = thisisyourid # please write int


@bot.command(aliases=['langs', 'languages', 'dil'])
async def diller(ctx):
    if str(ctx.message.channel.type) == "private":
        await ctx.message.channel.send("{0} ,Sorry I can't help you - Üzgünüm sana yardım edemem".format(ctx.message.author.mention),delete_after=5.0)
    else:
        dsc = ""

        channel = await ctx.message.author.create_dm()
        
        for i in googletrans.LANGUAGES:
            dsc += str(i) + " : " + googletrans.LANGUAGES[i] + "\n"

        embed = discord.Embed(title="**DİLLER - LANGUAGES**", colour=discord.Colour(0x3b8ff0), description=dsc)   # Please check this link (https://discordjs.guide/popular-topics/embeds.html#embed-preview)
        embed.set_footer(text="Made by DejaVu#4515")
        embed.set_author(name=ctx.message.author, icon_url=ctx.message.author.avatar_url)
        embed.set_thumbnail(url=ctx.message.author.avatar_url)
        
        await channel.send(embed=embed)
        await ctx.message.channel.send("{0} ,Sended dm to you for languages - Dm kutunuza diller gönderildi".format(ctx.message.author.mention),delete_after=5.0)
        await ctx.message.delete()

@bot.command(aliases=['aramak', 'arat', 'search'])
async def ara(ctx,dil=''):
    if str(ctx.message.channel.type) == "private":
        await ctx.message.channel.send("{0} ,Sorry I can't help you - Üzgünüm sana yardım edemem".format(ctx.message.author.mention),delete_after=5.0)
    else:
        if dil != '':
            dil = dil.lower()
            chk = 0
            for i,j in googletrans.LANGUAGES.items():
                if dil in j:
                    chk = 1
                    await ctx.message.channel.send("{0}, Language : {1}".format(ctx.message.author.mention,i),delete_after=15.0)
                    await ctx.message.delete()
                if chk == 0:
                    await ctx.message.channel.send("{0}, Ops! I didn't find your search - Aramanı bulamadım".format(ctx.message.author.mention),delete_after=5.0)
                    await ctx.message.delete()
                                
        else:
            await ctx.message.delete()
            await ctx.message.channel.send("{0}, Please type language - Lütfen dil yazınız".format(ctx.message.author.mention),delete_after=5.0)

@bot.command(aliases=['clean','silsüpür'])
async def purge(ctx,limit=''):
    if str(ctx.message.channel.type) == "private":
        await ctx.message.channel.send("{0} ,Sorry I can't help you - Üzgünüm sana yardım edemem".format(ctx.message.author.mention),delete_after=5.0)
    else:

        if ctx.message.author.guild_permissions.administrator:          # if you are administrator
            if limit != "":
                await ctx.message.channel.purge(limit=int(limit))
            else:
                await ctx.message.channel.purge()
        else:
            await ctx.message.channel.send("{0.author.mention}, şşş! kimse görmesin\nbu komudu kullanmamalısın".format(ctx.message),delete_after=5.0)
            await ctx.message.delete()

@bot.command()
async def sifir(ctx,mess=''):
    global chk
    if str(ctx.message.channel.type) == "private":
        await ctx.message.channel.send("{0} ,Sorry I can't help you - Üzgünüm sana yardım edemem".format(ctx.message.author.mention),delete_after=5.0)
    else:
        if str(ctx.message.author.id) == "596455467585110016":
            chk = 0
            await ctx.message.channel.send("{0.author.mention}, Done!".format(ctx.message),delete_after=5.0)
        else:
            await ctx.message.channel.send("{0.author.mention}, şşş! kimse görmesin\nbu komudu kullanmamalısın".format(ctx.message),delete_after=5.0)
            await ctx.message.delete()

@bot.command(aliases=['doküman','doc','word'])
async def docx(ctx,mess=''):
    if str(ctx.message.channel.type) == "private":
        await ctx.message.channel.send("{0} ,Sorry I can't help you - Üzgünüm sana yardım edemem".format(ctx.message.author.mention),delete_after=5.0)
    else:
        if mess != '':
            global chk
            if chk == 0:
                chk = 1
                document = Document()
                mess = mess.split("|")
                for i in range(len(mess)):
                    if mess[i] != "":
                        if mess[i][0] == "h":
                            document.add_heading(mess[i][1:len(mess[i])], 0)
                        elif mess[i][0] == "t":
                            document.add_paragraph(mess[i][1:len(mess[i])])
                        elif mess[i][0] == "m":
                            document.add_paragraph(
mess[i][1:len(mess[i])], style='List Bullet'
)
                        elif mess[i][0] == "l":
                            document.add_paragraph(
mess[i][1:len(mess[i])], style='List Number'
)
                
                document.save('example.docx')
                await ctx.message.channel.send("Done! - Oldu! {0}".format(ctx.message.author.mention),file=discord.File('example.docx'))
                chk = 0
            else:
                await ctx.message.channel.send("Please Try Again - Lütfen tekrar dene {0}".format(ctx.message.author.mention))
        else:
            await ctx.message.channel.send("Please fill in the required fields - Lütfen gerekli alanları doldurun {0}".format(ctx.message.author.mention))
        
    
    

@bot.command(aliases=['trans','cevir','translater'])
async def translate(ctx,cevir='',dest=''):

    dest = dest.lower()

    if str(ctx.message.channel.type) == "private":
        await ctx.message.channel.send("{0} ,Sorry I can't help you - Üzgünüm sana yardım edemem".format(ctx.message.author.mention),delete_after=5.0)
    else:

        cevir = cevir.replace("`","")
        
        if cevir != "":
            if dest != '':
                translator = Translator()
                print(dest)
                try:
                    dil = googletrans.LANGUAGES[translator.detect(cevir,dest=dest).lang]
                    dil2 = googletrans.LANGUAGES[dest]
                    cevirisi = translator.translate(cevir,dest=dest).text
                    await ctx.message.channel.send("{0.author.mention}, where did you find this language?\nTranslated language : {1}\nDestination Language: {3}\n```{2}```".format(ctx.message,dil,cevirisi,dil2))
                except:
                    await ctx.message.channel.send("{0.author.mention}, Ops!\nwrite enter valid language - geçerli bir dil girin".format(ctx.message),delete_after=5.0)
                    await ctx.message.delete()
            else:
                dest = 'en'
                translator = Translator()
                try:
                    dil = googletrans.LANGUAGES[translator.detect(cevir,dest=dest).lang]
                    dil2 = googletrans.LANGUAGES[dest]
                    cevirisi = translator.translate(cevir,dest=dest).text
                    await ctx.message.channel.send("{0.author.mention}, where did you find this language?\nTranslated language : {1}\nDestination Language: {3}\n```{2}```".format(ctx.message,dil,cevirisi,dil2))
                except:
                    await ctx.message.channel.send("{0.author.mention}, Ops!\nwrite enter valid language - geçerli bir dil girin".format(ctx.message),delete_after=5.0)
                    await ctx.message.delete()
        else:
            await ctx.message.channel.send("{0.author.mention}, Ops!\nwrite what you will translate - Neyi çevireceğini yaz...".format(ctx.message),delete_after=5.0)
            await ctx.message.delete()

@bot.command(aliases=['kelime', 'ogren'])
async def tdk(ctx,kelime=''):
    if str(ctx.message.channel.type) == "private":
        await ctx.message.channel.send("{0} ,Sorry I can't help you - Üzgünüm sana yardım edemem".format(ctx.message.author.mention),delete_after=5.0)
    else:
        if kelime != "":
            try:
                from tdk import tdk
                word = tdk.new_word(kelime)
                description = word.meaning()[0]
                embed = discord.Embed(title="**TDK ANLAMLARI : {0}**".format(kelime.lower()), colour=discord.Colour(0x3b8ff0), description=description)   # Please check this link (https://discordjs.guide/popular-topics/embeds.html#embed-preview)
                embed.set_footer(text="Made by DejaVu#4515")
                embed.set_author(name=ctx.message.author, icon_url=ctx.message.author.avatar_url)
                embed.set_thumbnail(url=ctx.message.author.avatar_url)
                await ctx.message.channel.send(embed=embed)
            except:
                await ctx.message.channel.send("Yazdığın kelimenin doğru olduğundan emin ol\n**Bulacağımdan eminim** {0}".format(ctx.message.author.mention))

        else:
            await ctx.message.channel.send("{0}, Ops! I didn't find your search - Aramanı bulamadım".format(ctx.message.author.mention),delete_after=5.0)
            await ctx.message.delete()
                        
@bot.command(aliases=['gazete', 'resmi','haber'])
async def resmigazete(ctx):
    if str(ctx.message.channel.type) == "private":
        await ctx.message.channel.send("{0} ,Sorry I can't help you - Üzgünüm sana yardım edemem".format(ctx.message.author.mention),delete_after=5.0)
    else:
        rsource = requests.get('https://www.resmigazete.gov.tr/')
        source = BeautifulSoup(rsource.content,"lxml")
        link = source.find("a",attrs={"id":"btnPdfGoruntule"}).get("href")
        yazi = source.find("div",attrs={"class":"html-subtitle"}).text
        aciklama2 = ""
        aciklama = source.find_all("a",attrs={"data-modal":"True"},limit=10)
        linka = source.find_all("a",attrs={"data-modal":"True"},limit=10)

        for i in range(len(aciklama)):
            aciklama2 += aciklama[i].text[0:len(aciklama[i].text)-15] + "[" + aciklama[i].text[len(aciklama[i].text)-15:len(aciklama[i].text)] + "](" + str(linka[i].get("href")) + ")\n\n"
        try:
            description = "[RESMİ GAZETE SON BASIM](" + str(link) + ")\n\n**" + yazi + "**\n\n" + aciklama2
            embed = discord.Embed(title="**RESMİ HABER**", colour=discord.Colour(0x3b8ff0), description=description)   # Please check this link (https://discordjs.guide/popular-topics/embeds.html#embed-preview)
            embed.set_footer(text="Made by DejaVu#4515")
            embed.set_author(name=ctx.message.author, icon_url=ctx.message.author.avatar_url)
            embed.set_thumbnail(url=ctx.message.author.avatar_url)
            await ctx.message.channel.send(embed=embed,delete_after=60.0)
        except:
            rsource = requests.get('https://www.resmigazete.gov.tr/')
            source = BeautifulSoup(rsource.content,"lxml")
            link = source.find("a",attrs={"id":"btnPdfGoruntule"}).get("href")
            yazi = source.find("div",attrs={"class":"html-subtitle"}).text
            aciklama2 = ""
            aciklama = source.find_all("a",attrs={"data-modal":"True"},limit=5)
            linka = source.find_all("a",attrs={"data-modal":"True"},limit=5)
            for i in range(len(aciklama)):
                aciklama2 += aciklama[i].text[0:len(aciklama[i].text)-15] + "[" + aciklama[i].text[len(aciklama[i].text)-15:len(aciklama[i].text)] + "](" + str(linka[i].get("href")) + ")\n\n"

            description = "[RESMİ GAZETE SON BASIM](" + str(link) + ")\n\n**" + yazi + "**\n\n" + aciklama2
            embed = discord.Embed(title="**RESMİ HABER**", colour=discord.Colour(0x3b8ff0), description=description)   # Please check this link (https://discordjs.guide/popular-topics/embeds.html#embed-preview)
            embed.set_footer(text="Made by DejaVu#4515 | Maddeler çok uzundu kısalttık :D İyi okumalar")
            embed.set_author(name=ctx.message.author, icon_url=ctx.message.author.avatar_url)
            embed.set_thumbnail(url=ctx.message.author.avatar_url)
            await ctx.message.channel.send(embed=embed,delete_after=60.0)

@bot.command(pass_context = True)
async def kick(ctx, member: discord.Member, reason=None):
    if ctx.message.author.guild_permissions.kick_members:
        try:
            await member.kick(reason=reason)
            a = await ctx.message.channel.send("{0}; Done, kicked **{1}!** - Bitti, **{1} atıldı!**".format(ctx.message.author.mention,str(member)))
            await a.add_reaction("\N{WHITE HEAVY CHECK MARK}")
        except:
            a = await ctx.message.channel.send("{0} ,You **can't** use that! - Bunu kullanamazsın! çünkü ondan **üstün değilsin**".format(ctx.message.author.mention))
            await a.add_reaction("🚫")
    else:
        try:
            await ctx.message.add_reaction("🚫")
        except:
            a = await ctx.message.channel.send("{0} ,You **can't** use that! - Bunu **kullanamazsın**!".format(ctx.message.author.mention))
            await a.add_reaction("🚫")


@bot.command(aliases=['help'])
async def yardim(ctx):

    if str(ctx.message.channel.type) == "private":
        await ctx.message.channel.send("{0} ,Sorry I can't help you - Üzgünüm sana yardım edemem".format(ctx.message.author.mention),delete_after=5.0)
    else:
        if ctx.message.content.startswith('>help') != True:
            try:
                utc_now = datetime.utcnow()
                utc = pytz.timezone('UTC')
                aware_date = utc.localize(utc_now)
                turkey = timezone('Europe/Istanbul')
                now_turkey = aware_date.astimezone(turkey)
                
                channel = await ctx.message.author.create_dm()
                embed = discord.Embed(title="Efsanevi Komutlarımız", colour=discord.Colour(0xd9b988), url="https://discordapp.com", description="Öncelikle, [botun](https://github.com/Deja-Vu1/Cevir-Bot) nasıl işlediğiniz merak ediyorsanız\n")
                embed.set_thumbnail(url=ctx.message.author.avatar_url)
                embed.set_author(name=str(ctx.message.author.name), icon_url=ctx.message.author.avatar_url)
                embed.set_footer(text="Made by DejaVu#4515 | {0}".format(str(now_turkey)[0:19]),icon_url="https://cdn.discordapp.com/avatars/596455467585110016/b96ac044a4382f62ad36637c6021ef80.png?size=128")

                embed.add_field(name="<:pencil:753955136824541215>", value="**>cevir** ``<text> <languag>``\nDilediğiniz en saçma şeyi bile çevirebilir...\n||öyle, değil mi?||", inline=True)
                embed.add_field(name="<:scroll:753956027983986688>", value="**>tdk** ``<word>``\nBilmemek değil öğrenmemek ayıp, _kelimeyi yaz anlamını bul!_", inline=True)
                embed.add_field(name="<:loudspeaker:753960791476142102>", value="**>haber**\nResmi gazeteyi önüne getir!\n_ilk 10 madde ||eğer çok uzunsa ilk 5||_", inline=True)
                embed.add_field(name="<:clipboard:753958902357295124>", value="**>docx** ``'tBaşlık|lsıralımadde|msırasızmadde|tparagraf'``\nSizin için not tutan bir asistanınız olsun ister miydiniz?\nNasıl kullanacağını öğren yazılarını online .docx'a çevir\n**-t :** 'text' paragraf ekler\n**-h :** 'heading' başlık ekler\n**-l :** 'list' sıralı madde ekler\n**-m :** 'madde' sırasız madde ekler ||hepsi ingilizce, bu neden türkçe :)||", inline=False)
                


                await channel.send(embed=embed)
                await ctx.message.channel.send("{0} ,Sended dm to you for commands - Dm kutunuza komutlar gönderildi".format(ctx.message.author.mention),delete_after=5.0)
                await ctx.message.add_reaction("\N{WHITE HEAVY CHECK MARK}")
            except:
                a = await ctx.message.channel.send("{0} ,I can't send dm to you - Sana dm gönderemiyorum".format(ctx.message.author.mention))
                await a.add_reaction("🚫")
        else:
            try:
                utc_now = datetime.utcnow()
                utc = pytz.timezone('UTC')
                aware_date = utc.localize(utc_now)
                turkey = timezone('Europe/Istanbul')
                now_turkey = aware_date.astimezone(turkey)
                
                channel = await ctx.message.author.create_dm()
                embed = discord.Embed(title="Legendary commands", colour=discord.Colour(0xd9b988), url="https://discordapp.com", description="First of all, if you are wondering how the [bot](https://github.com/Deja-Vu1/Cevir-Bot) works, look here\n")
                embed.set_thumbnail(url=ctx.message.author.avatar_url)
                embed.set_author(name=str(ctx.message.author.name), icon_url=ctx.message.author.avatar_url)
                embed.set_footer(text="Made by DejaVu#4515 | {0}".format(str(now_turkey)[0:19]),icon_url="https://cdn.discordapp.com/avatars/596455467585110016/b96ac044a4382f62ad36637c6021ef80.png?size=128")

                embed.add_field(name="<:pencil:753955136824541215>", value="**>cevir** ``<text> <language>``\nIt can translate even the most ridiculous thing you want...\n||Right?||", inline=True)
                embed.add_field(name="<:scroll:753956027983986688>", value="**>tdk** ``<word>``\nWrite the word and learn its meaning instantly. _isn't it very good?_", inline=True)
                embed.add_field(name="<:loudspeaker:753960791476142102>", value="**>haber**\nBring the first 10 articles of the official newspaper in front of you!\n||if 10 items are too long, you can read 5 items||", inline=True)
                embed.add_field(name="<:clipboard:753958902357295124>", value="**>docx** ``'htitle|ladditemwithnumbers|madditemwithoutnumbers|tparagraph'``\nWould you like to have a note-taking assistant for you?\nLearn how to use convert your articles for write .docx on online\n**-t :** 'text' add paragraph\n**-h :** 'heading' add title\n**-l :** 'list' add item with numbers\n**-m :** 'madde' add item without numbers", inline=False)
                


                await channel.send(embed=embed)
                await ctx.message.channel.send("{0} ,Sended dm to you for commands - Dm kutunuza komutlar gönderildi".format(ctx.message.author.mention),delete_after=5.0)
                await ctx.message.add_reaction("\N{WHITE HEAVY CHECK MARK}")
            except:
                a = await ctx.message.channel.send("{0} ,I can't send dm to you - Sana dm gönderemiyorum".format(ctx.message.author.mention))
                await a.add_reaction("🚫")


@bot.event
async def on_ready():
    await bot.change_presence(
    activity=discord.Game(
    name=">yardim & >help | 🌐 " + str(len(bot.guilds)) + " servers | Made by DejaVu#4515\n"))

@bot.event
async def on_guild_join(guild):
    global adminid
    user = bot.get_user(adminid)
    embed = discord.Embed(title="**NEW SERVER**", colour=discord.Colour(0x4aff00), description="\n**Members:** " + str(len(guild.members)) + "\n**Banner:** [click](" + str(guild.banner_url) + ")\n**Owner:** " + str(guild.owner) + "\n**Server Id:** " + str(guild.id))   # Please check this link (https://discordjs.guide/popular-topics/embeds.html#embed-preview)
    embed.set_footer(text="Information Service")
    embed.set_author(name=guild, icon_url=guild.icon_url)
    embed.set_thumbnail(url=guild.icon_url)
    await user.send(embed=embed)
    await bot.change_presence(
        activity=discord.Game(
            name=">yardim & >help | 🌐 " + str(len(bot.guilds)) + " servers | Made by DejaVu#4515\n"))

@bot.event
async def on_guild_remove(guild):
    global adminid
    user = bot.get_user(adminid)
    embed = discord.Embed(title="**LEAVED**", colour=discord.Colour(0xd0021b), description="\n**Members:** " + str(len(guild.members)) + "\n**Banner:** [click](" + str(guild.banner_url) + ")\n**Owner:** " + str(guild.owner) + "\n**Server Id:** " + str(guild.id))   # Please check this link (https://discordjs.guide/popular-topics/embeds.html#embed-preview)
    embed.set_footer(text="Information Service")
    embed.set_author(name=guild, icon_url=guild.icon_url)
    embed.set_thumbnail(url=guild.icon_url)
    await user.send(embed=embed)
    await bot.change_presence(
        activity=discord.Game(
            name=">yardim & >help | 🌐 " + str(len(bot.guilds)) + " servers | Made by DejaVu#4515\n"))

bot.run('TOKEN')
